import docx
import fitz  # PyMuPDF
from langchain.schema import Document

def process_docx(file_path):
    try:
        doc = docx.Document(file_path)
        elements = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue  # ignorar párrafos vacíos

            style_name = para.style.name if para.style else "Normal"

            # Si es un título → convertirlo en Markdown heading
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    level = 1  # fallback por si algo raro viene en estilos
                elements.append(f"{'#' * level} {text}")
            else:
                elements.append(text)

        # Unir con saltos de línea dobles para respetar formato Markdown
        structured_text = "\n\n".join(elements)

        return [Document(
            page_content=structured_text,
            metadata={"source": str(file_path), "type": "docx"}
        )]

    except Exception as e:
        print(f"Error procesando DOCX {file_path}: {e}")
        return []

def process_md(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        # Normalizar: quitar espacios sobrantes y asegurar saltos de línea consistentes
        text = "\n".join(line.rstrip() for line in text.splitlines())
        
        return [Document(
            page_content=text.strip(),
            metadata={"source": str(file_path), "type": "md"})
        ]
    except Exception as e:
        print(f"Error procesando MD {file_path}: {e}")
        return []
